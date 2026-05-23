"""
documentation.py

Onglet Documentation — Keepinbot
==================================
Interface du Module 2 — Génération de documentation.

Workflow utilisateur :
1. Uploader plusieurs fichiers (Word, PPT, mail, texte)
2. Donner un titre au document à générer
3. Cliquer sur Générer
4. Relire le document produit
5. Valider pour l'indexer dans la base RAG — ou rejeter

Principe de validation humaine :
Aucun document généré n'est automatiquement indexé.
L'utilisateur doit explicitement valider avant indexation.
"""

import streamlit as st
import os
import tempfile
from docx import Document as DocxDocument
from app.core.generator import parse_file, generate_document
from app.core.rag import load_documents, chunk_documents, build_vectorstore
from app.core.config import CORPUS_PATH


def save_as_docx(content: str, title: str) -> bytes:
    """
    Convertit un document Markdown en fichier Word (.docx) téléchargeable.

    Principe :
    Le LLM génère du Markdown — on le convertit en Word pour que
    l'utilisateur puisse le relire et le modifier dans son traitement
    de texte habituel avant de le valider.

    Paramètres :
    - content : texte Markdown généré par le LLM
    - title   : titre du document, utilisé comme titre Word

    Retourne :
    - contenu binaire du fichier .docx (bytes) prêt à télécharger
    """
    # Nettoyage des balises Markdown que le LLM ajoute parfois
    # autour du contenu (```markdown ... ```)
    content = content.strip()
    if content.startswith("```markdown"):
        content = content[len("```markdown"):].strip()
    if content.startswith("```"):
        content = content[3:].strip()
    if content.endswith("```"):
        content = content[:-3].strip()

    doc = DocxDocument()
    doc.add_heading(title, 0)

    # Conversion simple : chaque ligne devient un paragraphe
    # Les lignes commençant par ## deviennent des titres de niveau 1
    # Les lignes commençant par ### deviennent des titres de niveau 2
    for line in content.split("\n"):
        line = line.strip()
        if not line or line == "---":
            continue
        if line.startswith("### "):
            doc.add_heading(line[4:], level=2)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=1)
        elif line.startswith("# "):
            # Ignore le titre principal s'il est identique au titre du document
            heading_text = line[2:].strip()
            if heading_text.lower() != title.lower():
                doc.add_heading(heading_text, level=0)
        else:
            line = line.replace("**", "").replace("*", "")
            doc.add_paragraph(line)

    # Sauvegarde en mémoire (pas sur disque) pour le téléchargement
    import io
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()


def index_document(content: str, title: str) -> bool:
    """
    Indexe un document validé dans la base RAG.
    Sauvegarde le contenu dans data/corpus/ puis reconstruit le vectorstore.

    Paramètres :
    - content : texte du document à indexer
    - title   : titre utilisé pour nommer le fichier

    Retourne :
    - True si l'indexation a réussi, False sinon

    Note : le fichier est sauvegardé avec le préfixe "generated_"
    pour le distinguer des documents sources originaux.
    """
    try:
        # Nettoyage du titre pour en faire un nom de fichier valide
        safe_title = title.lower().replace(" ", "_").replace("/", "-")
        filename = f"generated_{safe_title}.txt"
        filepath = os.path.join(CORPUS_PATH, filename)

        with open(filepath, "w", encoding="utf-8") as f:
            f.write(f"# {title}\n\n{content}")

        # Reconstruction du vectorstore avec le nouveau document
        docs = load_documents(CORPUS_PATH)
        chunks = chunk_documents(docs)
        build_vectorstore(chunks)

        return True
    except Exception as e:
        st.error(f"Erreur lors de l'indexation : {e}")
        return False


def render_documentation():
    """
    Rendu de l'onglet Documentation.
    Gère le workflow en 3 étapes via st.session_state :
    - Étape 1 : upload et configuration
    - Étape 2 : aperçu du document généré
    - Étape 3 : confirmation d'indexation
    """

    st.subheader("Génération de documentation")
    st.caption("Uploadez des fichiers hétérogènes — le système génère un document structuré.")

    # ── Initialisation de l'état de session ───────────────────────────────────
    if "doc_generated" not in st.session_state:
        st.session_state.doc_generated = None  # document généré en attente
    if "doc_title" not in st.session_state:
        st.session_state.doc_title = ""

    # ── Étape 1 : Upload et configuration ────────────────────────────────────
    st.markdown("### Étape 1 — Uploader les fichiers sources")

    uploaded_files = st.file_uploader(
        "Sélectionnez un ou plusieurs fichiers",
        type=["txt", "docx", "pptx", "eml"],
        accept_multiple_files=True,  # upload multiple activé
        help="Formats acceptés : .txt, .docx, .pptx, .eml"
    )

    doc_title = st.text_input(
        "Titre du document à générer",
        placeholder="Ex : Synthèse projet Alpha, Procédure onboarding...",
        value=st.session_state.doc_title
    )

    if uploaded_files and doc_title:
        if st.button("Générer le document", type="primary"):
            st.session_state.doc_title = doc_title

            # Sauvegarde temporaire des fichiers uploadés pour les parser
            parsed = []
            with tempfile.TemporaryDirectory() as tmpdir:
                for uploaded in uploaded_files:
                    tmp_path = os.path.join(tmpdir, uploaded.name)
                    with open(tmp_path, "wb") as f:
                        f.write(uploaded.getbuffer())
                    result = parse_file(tmp_path)
                    if result["error"]:
                        st.warning(f"Impossible de lire {uploaded.name} : {result['error']}")
                    else:
                        parsed.append(result)
                        st.caption(f"✓ {uploaded.name} — {len(result['content'])} caractères")

            if parsed:
                with st.spinner("Génération en cours... (peut prendre 1-2 minutes en mode local)"):
                    generated = generate_document(parsed, doc_title)

                if generated.startswith("Erreur"):
                    st.error(generated)
                    st.info("Conseil : bascule en mode cloud dans le .env pour la génération de documents longs.")
                else:
                    st.session_state.doc_generated = generated
                    st.rerun()
            else:
                st.error("Aucun fichier n'a pu être parsé.")

    elif uploaded_files and not doc_title:
        st.info("Donnez un titre au document avant de générer.")

    # ── Étape 2 : Aperçu et validation ───────────────────────────────────────
    if st.session_state.doc_generated:
        st.divider()
        st.markdown("### Étape 2 — Relire et valider")
        st.caption("⚠️ Ce document est un brouillon généré automatiquement. Relisez-le avant validation.")

        # Aperçu du document généré
        with st.expander("Aperçu du document généré", expanded=True):
            st.markdown(st.session_state.doc_generated)

        # Boutons d'action
        col1, col2, col3 = st.columns(3)

        with col1:
            # Téléchargement Word
            docx_bytes = save_as_docx(
                st.session_state.doc_generated,
                st.session_state.doc_title
            )
            st.download_button(
                label="Télécharger en Word",
                data=docx_bytes,
                file_name=f"{st.session_state.doc_title.replace(' ', '_')}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

        with col2:
            # Validation et indexation dans la base RAG
            if st.button("Valider et indexer", type="primary"):
                with st.spinner("Indexation en cours..."):
                    success = index_document(
                        st.session_state.doc_generated,
                        st.session_state.doc_title
                    )
                if success:
                    st.success("Document indexé dans la base RAG. Il est maintenant interrogeable via l'onglet Assistant.")
                    st.session_state.doc_generated = None
                    st.session_state.doc_title = ""
                    st.rerun()

        with col3:
            # Rejet — on efface le document généré sans indexer
            if st.button("Rejeter"):
                st.session_state.doc_generated = None
                st.session_state.doc_title = ""
                st.info("Document rejeté — non indexé.")
                st.rerun()